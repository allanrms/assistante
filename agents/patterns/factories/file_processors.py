"""
Processadores de arquivos para extrair contexto de diferentes tipos de arquivo
Suporta extração via IA (GPT-4o) para PDFs, planilhas e imagens
"""
import os
import mimetypes
import traceback
from typing import Optional, Dict, Any
import base64
import json
import csv
import io
import re
import tempfile

# Imports de terceiros
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import PyPDF4 as PyPDF2
except ImportError:
    try:
        import PyPDF2
    except ImportError:
        PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

# LangChain e LLM
from agents.patterns.factories.llm_factory import LLMFactory
from langchain_core.messages import HumanMessage


class BaseFileProcessor:
    """Classe base para processadores de arquivo"""
    
    def __init__(self):
        self.supported_extensions = []
        self.max_file_size = 10 * 1024 * 1024  # 10MB por padrão
    
    def can_process(self, file_path: str) -> bool:
        """Verifica se pode processar o arquivo baseado na extensão"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.supported_extensions
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto do arquivo. Deve ser implementado pelas subclasses."""
        raise NotImplementedError
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Retorna informações básicas do arquivo"""
        stat = os.stat(file_path)
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {
            'size': stat.st_size,
            'mime_type': mime_type,
            'extension': os.path.splitext(file_path)[1].lower()
        }


class TextFileProcessor(BaseFileProcessor):
    """Processador para arquivos de texto simples"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt', '.md', '.csv', '.json', '.html', '.xml']
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto de arquivos texto"""
        try:
            # Tentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                    # Se conseguiu ler, processa baseado na extensão
                    ext = os.path.splitext(file_path)[1].lower()
                    
                    if ext == '.json':
                        return self._process_json(content)
                    elif ext == '.csv':
                        return self._process_csv(content)
                    elif ext == '.html' or ext == '.xml':
                        return self._process_html(content)
                    else:
                        return content.strip()
                        
                except UnicodeDecodeError:
                    continue
                    
            return "Erro: Não foi possível decodificar o arquivo"
            
        except Exception as e:
            return f"Erro ao processar arquivo: {str(e)}"
    
    def _process_json(self, content: str) -> str:
        """Formata JSON para melhor legibilidade"""
        try:
            data = json.loads(content)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except:
            return content
    
    def _process_csv(self, content: str) -> str:
        """Converte CSV em formato legível"""
        try:
            # Detectar delimitador
            sample = content[:1024]
            sniffer = csv.Sniffer()
            delimiter = sniffer.sniff(sample).delimiter
            
            # Converter para formato tabular
            reader = csv.reader(io.StringIO(content), delimiter=delimiter)
            rows = list(reader)
            
            if not rows:
                return content
                
            # Formatar como tabela
            formatted_rows = []
            for i, row in enumerate(rows[:100]):  # Limitar a 100 linhas
                if i == 0:
                    formatted_rows.append("Cabeçalhos: " + " | ".join(row))
                    formatted_rows.append("-" * 50)
                else:
                    formatted_rows.append(" | ".join(row))
                    
            if len(rows) > 100:
                formatted_rows.append(f"... e mais {len(rows) - 100} linhas")
                
            return "\n".join(formatted_rows)
            
        except:
            return content
    
    def _process_html(self, content: str) -> str:
        """Extrai texto de HTML"""
        try:
            if BeautifulSoup is None:
                return content
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remover scripts e estilos
            for script in soup(["script", "style"]):
                script.decompose()
                
            # Extrair texto
            text = soup.get_text()
            
            # Limpar linhas vazias excessivas
            lines = [line.strip() for line in text.splitlines()]
            lines = [line for line in lines if line]

            return '\n'.join(lines)

        except:
            return content


class PDFFileProcessor(BaseFileProcessor):
    """Processador para arquivos PDF"""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']

    def extract_text(self, file_path: str) -> str:
        """Extrai texto de PDFs"""
        try:
            if PyPDF2 is None:
                return "PyPDF2 ou PyPDF4 não está instalado. Instale com: pip install PyPDF4"

            text_content = []

            with open(file_path, 'rb') as file:
                try:
                    # Tentar com PdfReader (versão mais nova)
                    if hasattr(PyPDF2, 'PdfReader'):
                        pdf_reader = PyPDF2.PdfReader(file)
                        pages = pdf_reader.pages
                    else:
                        # Fallback para versão antiga
                        pdf_reader = PyPDF2.PdfFileReader(file)
                        pages = [pdf_reader.getPage(i) for i in range(pdf_reader.numPages)]

                    for page_num, page in enumerate(pages):
                        try:
                            # Tentar extract_text() primeiro (versão nova)
                            if hasattr(page, 'extract_text'):
                                page_text = page.extract_text()
                            else:
                                # Fallback para versão antiga
                                page_text = page.extractText()

                            if page_text and page_text.strip():
                                # Limpar texto extraído
                                cleaned_text = self._clean_pdf_text(page_text)
                                if cleaned_text:
                                    text_content.append(f"=== Página {page_num + 1} ===")
                                    text_content.append(cleaned_text)
                                    text_content.append("")
                        except Exception as e:
                            text_content.append(f"[Erro na página {page_num + 1}: {str(e)}]")

                except Exception as e:
                    return f"Erro ao abrir PDF: {str(e)}"

            if text_content:
                result = '\n'.join(text_content)
                return result if result.strip() else "PDF processado mas não foi possível extrair texto legível"
            else:
                return "Não foi possível extrair texto do PDF"

        except Exception as e:
            return f"Erro ao processar PDF: {str(e)}"

    def _clean_pdf_text(self, text: str) -> str:
        """Limpa texto extraído de PDF removendo caracteres especiais e formatação estranha"""
        # Remover caracteres de controle e especiais
        # Manter apenas letras, números, pontuação comum e espaços
        text = re.sub(r'[^\w\s\.,;:!?@#$%&*()\-+=\[\]{}|\\/"\'<>áàâãéèêíïóôõöúçñÁÀÂÃÉÈÊÍÏÓÔÕÖÚÇÑ]', ' ', text)

        # Remover múltiplas quebras de linha consecutivas
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        # Remover múltiplos espaços
        text = re.sub(r' +', ' ', text)

        # Remover espaços no início e fim de cada linha
        lines = [line.strip() for line in text.split('\n')]

        # Remover linhas vazias
        lines = [line for line in lines if line]

        # Juntar linhas muito curtas (provavelmente cabeçalhos quebrados)
        cleaned_lines = []
        i = 0
        while i < len(lines):
            line = lines[i]

            # Se a linha tem menos de 40 caracteres e a próxima existe
            # provavelmente é uma quebra de linha indevida
            if i + 1 < len(lines) and len(line) < 40 and not line.endswith('.'):
                # Juntar com a próxima linha
                line = line + ' ' + lines[i + 1]
                i += 2
            else:
                i += 1

            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)


class DocxFileProcessor(BaseFileProcessor):
    """Processador para arquivos Word DOCX"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    def extract_text(self, file_path: str) -> str:
        """Extrai texto de arquivos DOCX"""
        try:
            if docx is None:
                return "python-docx não está instalado. Instale com: pip install python-docx"

            doc = docx.Document(file_path)
            text_content = []
            
            # Extrair texto dos parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extrair texto de tabelas
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("--- Tabela ---")
                    text_content.extend(table_text)
                    text_content.append("")


            return '\n'.join(text_content) if text_content else "Documento vazio"

        except Exception as e:
            return f"Erro ao processar DOCX: {str(e)}"


class ImageFileProcessor(BaseFileProcessor):
    """Processador para arquivos de imagem"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff']
        self.max_file_size = 20 * 1024 * 1024  # 20MB para imagens
    
    def extract_text(self, file_path: str) -> str:
        """
        Para imagens, não extraímos texto diretamente.
        Retornamos metadados básicos da imagem.
        """
        try:
            if Image is None:
                return "PIL (Pillow) não está instalado. Instale com: pip install Pillow"

            # Abrir imagem e extrair metadados
            with Image.open(file_path) as img:
                # Informações básicas da imagem
                width, height = img.size
                mode = img.mode
                format_name = img.format or "Unknown"
                
                # Tamanho do arquivo
                file_size = os.path.getsize(file_path)
                size_mb = file_size / (1024 * 1024)
                
                # Construir descrição da imagem
                description = f"""Imagem: {os.path.basename(file_path)}
Formato: {format_name}
Dimensões: {width} x {height} pixels
Modo de cor: {mode}
Tamanho: {size_mb:.2f} MB

Nota: Esta é uma imagem que pode ser enviada diretamente ao usuário via WhatsApp.
Para análise visual do conteúdo da imagem, ela será processada pela IA Vision quando necessário."""
                
                # Verificar se há metadados EXIF (opcional)
                if hasattr(img, '_getexif') and img._getexif():
                    description += "\n\nMetadados EXIF disponíveis."


                return description

        except Exception as e:
            return f"Erro ao processar imagem: {str(e)}"


class ExcelFileProcessor(BaseFileProcessor):
    """Processador para arquivos Excel (XLS, XLSX)"""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.xls', '.xlsx']

    def extract_text(self, file_path: str) -> str:
        """Extrai texto de arquivos Excel"""
        try:
            if pd is None:
                return "pandas e openpyxl não estão instalados. Instale com: pip install pandas openpyxl"

            # Ler todas as planilhas
            excel_file = pd.ExcelFile(file_path)
            all_text = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                all_text.append(f"=== Planilha: {sheet_name} ===")
                all_text.append("")

                # Converter DataFrame para texto formatado
                # Limitar a 1000 linhas para não sobrecarregar
                if len(df) > 1000:
                    df = df.head(1000)
                    all_text.append(f"[Mostrando primeiras 1000 de {len(df)} linhas]")
                    all_text.append("")

                # Converter para string mantendo formatação tabular
                table_text = df.to_string(index=False)
                all_text.append(table_text)
                all_text.append("")

            return '\n'.join(all_text)

        except Exception as e:
            return f"Erro ao processar Excel: {str(e)}"


class AIFileProcessor:
    """Processador que usa IA para extrair conteúdo de arquivos complexos"""

    def __init__(self, agent=None):
        """
        Args:
            agent: Instância do Agent para usar o LLM configurado (opcional)
        """
        self.agent = agent
        self.supported_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.webp', '.xls', '.xlsx']

    def can_process(self, file_path: str) -> bool:
        """Verifica se pode processar o arquivo"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.supported_extensions

    def extract_with_ai(self, file_path: str, use_vision: bool = False) -> Dict[str, Any]:
        """
        Extrai conteúdo usando LLM configurado via LLMFactory.

        Args:
            file_path: Caminho do arquivo
            use_vision: Se True, usa vision para processar como imagem

        Returns:
            Dict com 'success', 'extracted_text' ou 'error'
        """
        try:
            # Criar LLMFactory uma única vez e reutilizar
            llm_factory = LLMFactory(agent=self.agent)
            llm = llm_factory.llm

            if use_vision:
                # Processar como imagem usando Vision
                return self._extract_with_vision(file_path, llm)
            else:
                # Processar PDF convertendo para imagens
                return self._extract_pdf_with_vision(file_path, llm)

        except Exception as e:
            return {
                'success': False,
                'error': f'Erro ao processar com IA: {str(e)}'
            }

    def _extract_with_vision(self, file_path: str, llm) -> Dict[str, Any]:
        """Extrai conteúdo de imagem usando LLM com Vision"""
        try:

            # Ler arquivo e converter para base64
            with open(file_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')

            # Determinar tipo MIME
            ext = os.path.splitext(file_path)[1].lower()
            mime_types = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.webp': 'image/webp'
            }
            mime_type = mime_types.get(ext, 'image/jpeg')

            # Criar mensagem com imagem usando LangChain
            message = HumanMessage(
                content=[
                    {
                        "type": "text",
                        "text": """Você é um sistema de OCR profissional. Extraia TODO o texto desta imagem seguindo estas regras:

1. Leia o texto NA ORDEM CORRETA (da esquerda para direita, de cima para baixo)
2. Preserve a formatação original (parágrafos, quebras de linha, espaçamento)
3. Se houver tabelas, formate-as usando | (pipe) como separador
4. Se houver listas, mantenha os marcadores (bullets/números)
5. Mantenha títulos e subtítulos claramente identificados
6. NÃO adicione comentários, explicações ou observações
7. NÃO traduza ou modifique o texto
8. Retorne APENAS o texto extraído, exatamente como aparece

Se não houver texto na imagem, retorne apenas: "[Sem texto detectado]"
"""
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{image_data}",
                            "detail": "high"
                        }
                    }
                ]
            )

            # Invocar LLM
            response = llm.invoke([message])
            extracted_text = response.content.strip()

            # Verificar se o texto foi extraído com sucesso
            if not extracted_text or extracted_text == "[Sem texto detectado]":
                return {
                    'success': False,
                    'error': 'Nenhum texto detectado na imagem'
                }

            return {
                'success': True,
                'extracted_text': extracted_text,
                'method': 'llm-vision'
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'Erro no Vision: {str(e)}'
            }

    def _extract_pdf_with_vision(self, file_path: str, llm) -> Dict[str, Any]:
        """Converte PDF para imagens e extrai conteúdo com LLM Vision"""
        try:
            if convert_from_path is None:
                return {
                    'success': False,
                    'error': 'pdf2image não está instalado. Instale com: pip install pdf2image'
                }

            print(f"📄 Convertendo PDF para imagens em alta resolução...")
            # Converter PDF para imagens em ALTA RESOLUÇÃO para melhor OCR
            # DPI 300 garante qualidade profissional para OCR
            images = convert_from_path(
                file_path,
                dpi=300,  # Alta resolução para OCR de qualidade
                first_page=1,
                last_page=10,
                fmt='png'  # PNG preserva qualidade
            )
            print(f"✅ {len(images)} página(s) convertida(s) em alta resolução")

            all_text = []

            for i, image in enumerate(images, 1):
                # Salvar imagem temporária
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                    image.save(temp_file.name, 'PNG')
                    temp_path = temp_file.name

                try:
                    # Extrair texto da imagem usando LLM Vision (reutiliza mesma instância)
                    result = self._extract_with_vision(temp_path, llm)

                    if result['success']:
                        all_text.append(f"=== Página {i} ===")
                        all_text.append(result['extracted_text'])
                        all_text.append("")
                    else:
                        all_text.append(f"[Erro na página {i}: {result.get('error', 'Desconhecido')}]")

                finally:
                    # Remover arquivo temporário
                    if os.path.exists(temp_path):
                        os.remove(temp_path)

            if len(images) > 10:
                all_text.append(f"\n[Documento possui {len(images)} páginas. Processadas apenas as primeiras 10.]")

            return {
                'success': True,
                'extracted_text': '\n'.join(all_text),
                'method': 'pdf2image + llm-vision'
            }

        except Exception as e:
            traceback.print_exc()
            return {
                'success': False,
                'error': f'Erro ao processar PDF com Vision: {str(e)}'
            }


class FileProcessorFactory:
    """Factory para criar processadores baseado no tipo de arquivo"""

    def __init__(self, agent=None):
        """
        Args:
            agent: Instância do Agent para usar o LLM configurado (opcional)
        """
        self.agent = agent
        self.processors = {
            'text': TextFileProcessor(),
            'pdf': PDFFileProcessor(),
            'docx': DocxFileProcessor(),
            'image': ImageFileProcessor(),
            'excel': ExcelFileProcessor(),
        }
        self.ai_processor = AIFileProcessor(agent=agent)
    
    def get_processor(self, file_path: str) -> Optional[BaseFileProcessor]:
        """Retorna o processador apropriado para o arquivo"""
        for processor in self.processors.values():
            if processor.can_process(file_path):
                return processor
        return None
    
    def process_file(self, file_path: str, use_ai: bool = True) -> Dict[str, Any]:
        """
        Processa arquivo e retorna informações + conteúdo extraído.

        Args:
            file_path: Caminho do arquivo
            use_ai: Se True, usa IA (GPT-4o) como fallback ou para arquivos complexos

        Returns:
            Dict com 'success', 'extracted_text', 'error', 'file_info', 'method'
        """
        processor = self.get_processor(file_path)

        if not processor:
            return {
                'success': False,
                'error': 'Tipo de arquivo não suportado',
                'extracted_text': '',
                'file_info': {},
                'method': 'none'
            }

        try:
            # Verificar tamanho do arquivo
            file_info = processor.get_file_info(file_path)
            if file_info['size'] > processor.max_file_size:
                return {
                    'success': False,
                    'error': f'Arquivo muito grande. Máximo: {processor.max_file_size / (1024*1024):.1f}MB',
                    'extracted_text': '',
                    'file_info': file_info,
                    'method': 'none'
                }

            # 1. Tentar extração tradicional primeiro
            extracted_text = processor.extract_text(file_path)

            # 2. Verificar se a extração foi bem-sucedida
            # Para PDFs e imagens, SEMPRE usar IA para melhor OCR
            should_use_ai = False

            if use_ai and self.ai_processor.can_process(file_path):
                ext = os.path.splitext(file_path)[1].lower()

                # SEMPRE usar IA para PDFs e imagens (melhor OCR)
                if ext in ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.webp']:
                    should_use_ai = True
                    print(f"🤖 Usando IA para extração de {ext.upper()} (OCR de alta qualidade)...")
                # Usar IA para planilhas se a extração tradicional falhar
                elif ext in ['.xls', '.xlsx']:
                    if len(extracted_text.strip()) < 50 or 'Erro' in extracted_text:
                        should_use_ai = True
                        print(f"⚠️ Extração tradicional insuficiente, usando IA...")
                # Usar IA se texto muito curto ou com erro
                elif len(extracted_text.strip()) < 50:
                    should_use_ai = True
                    print(f"⚠️ Texto extraído muito curto ({len(extracted_text)} chars), usando IA...")
                elif 'Erro' in extracted_text or 'não foi possível' in extracted_text:
                    should_use_ai = True
                    print(f"⚠️ Erro na extração tradicional, usando IA...")

            # 3. Usar IA se necessário
            if should_use_ai:
                # Obter nome do modelo para print correto
                model_name = f"{self.agent.name} ({self.agent.model})" if self.agent else "GPT-4o"
                print(f"🤖 Usando IA para extração: {model_name}...")
                ext = os.path.splitext(file_path)[1].lower()

                if ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
                    # Processar imagem diretamente com Vision
                    ai_result = self.ai_processor.extract_with_ai(file_path, use_vision=True)
                elif ext == '.pdf':
                    # Processar PDF convertendo para imagens
                    ai_result = self.ai_processor.extract_with_ai(file_path, use_vision=False)
                else:
                    ai_result = {'success': False, 'error': 'Tipo não suportado pela IA'}

                if ai_result['success']:
                    return {
                        'success': True,
                        'error': None,
                        'extracted_text': ai_result['extracted_text'],
                        'file_info': file_info,
                        'method': ai_result.get('method', 'ai')
                    }
                else:
                    # Se IA falhar, retornar o texto extraído tradicionalmente
                    print(f"⚠️ IA falhou: {ai_result.get('error')}. Usando extração tradicional.")

            # 4. Retornar resultado da extração tradicional
            return {
                'success': True,
                'error': None,
                'extracted_text': extracted_text,
                'file_info': file_info,
                'method': 'traditional'
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'Erro durante processamento: {str(e)}',
                'extracted_text': '',
                'file_info': {},
                'method': 'error'
            }
    
    def get_supported_extensions(self) -> list:
        """Retorna lista de extensões suportadas"""
        extensions = []
        for processor in self.processors.values():
            extensions.extend(processor.supported_extensions)
        return sorted(list(set(extensions)))